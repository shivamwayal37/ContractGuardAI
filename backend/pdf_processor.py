import PyPDF2
import docx
import logging
from typing import Optional

logger = logging.getLogger(__name__)

class PDFProcessor:
    """
    Handles extraction of text from various document formats
    Supports PDF, DOCX, and TXT files
    """
    
    def extract_text_from_pdf(self, filepath: str) -> str:
        """
        Extract text from a PDF file
        
        Args:
            filepath: Path to the PDF file
        
        Returns:
            Extracted text as string
        """
        try:
            text = ""
            with open(filepath, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                
                logger.info(f"Processing PDF with {num_pages} pages")
                
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    
                    # Add page separator for context
                    text += f"\n--- Page {page_num + 1} ---\n"
                    text += page_text
                
                # Clean up the text
                text = self._clean_text(text)
                
                logger.info(f"Extracted {len(text)} characters from PDF")
                
                return text
        
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            raise Exception(f"Failed to process PDF: {str(e)}")
    
    def extract_text_from_docx(self, filepath: str) -> str:
        """
        Extract text from a DOCX file
        
        Args:
            filepath: Path to the DOCX file
        
        Returns:
            Extracted text as string
        """
        try:
            doc = docx.Document(filepath)
            text = ""
            
            # Extract paragraphs
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract tables if any
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            # Clean up the text
            text = self._clean_text(text)
            
            logger.info(f"Extracted {len(text)} characters from DOCX")
            
            return text
        
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise Exception(f"Failed to process DOCX: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """
        Clean extracted text
        - Remove excessive whitespace
        - Normalize line breaks
        - Remove special characters that might confuse AI
        
        Args:
            text: Raw extracted text
        
        Returns:
            Cleaned text
        """
        # Replace multiple spaces with single space
        text = ' '.join(text.split())
        
        # Normalize line breaks
        text = text.replace('\r\n', '\n')
        text = text.replace('\r', '\n')
        
        # Remove multiple consecutive line breaks
        while '\n\n\n' in text:
            text = text.replace('\n\n\n', '\n\n')
        
        # Remove null bytes and other problematic characters
        text = text.replace('\x00', '')
        text = text.replace('\ufffd', '')
        
        return text.strip()
    
    def validate_text(self, text: str, min_length: int = 100) -> bool:
        """
        Validate that extracted text is usable
        
        Args:
            text: Extracted text
            min_length: Minimum acceptable length
        
        Returns:
            True if text is valid, False otherwise
        """
        if not text or len(text.strip()) < min_length:
            return False
        
        # Check if text is mostly readable (not just garbage characters)
        alphanumeric_count = sum(c.isalnum() or c.isspace() for c in text)
        if alphanumeric_count / len(text) < 0.5:
            logger.warning("Text appears to contain mostly non-readable characters")
            return False
        
        return True